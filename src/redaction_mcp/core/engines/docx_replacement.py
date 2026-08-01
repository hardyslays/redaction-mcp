"""python-docx implementation of permanent text data replacement."""

from __future__ import annotations

import base64
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument

from redaction_mcp.core.engines.docx_redaction import _paragraphs
from redaction_mcp.core.models.document import Document
from redaction_mcp.core.models.replacement import ReplacementTarget
from redaction_mcp.core.services.replacement_text import replacement_text
from redaction_mcp.core.services.text_matching import compile_text_pattern


def _replace(paragraph: object, pattern: re.Pattern[str], target: ReplacementTarget) -> None:
    text, count = pattern.subn(lambda match: replacement_text(match.group(), target), paragraph.text)
    if count:
        # Rebuilding the paragraph removes the original runs from the DOCX package.
        paragraph.clear()
        paragraph.add_run(text)


def replace_docx_document(document: Document, targets: list[ReplacementTarget], *, data: bytes | None = None) -> Document:
    """Replace exact text in body, table, header, and footer paragraphs."""
    if data is None and document.base64data is None:
        raise ValueError("DOCX engine requires in-memory document base64data")
    source = DocxDocument(BytesIO(data if data is not None else document.decoded_bytes()))
    paragraphs = _paragraphs(source)
    for target in targets:
        if target.pages is not None:
            raise ValueError("DOCX data replacement does not support page-restricted targets")
        pattern = compile_text_pattern(
            target.values, ignore_case=target.ignore_case, partial_match=target.partial_match
        )
        for paragraph in paragraphs:
            _replace(paragraph, pattern, target)
    output = BytesIO()
    source.save(output)
    filename = Path(document.filename or "document.docx")
    return Document(
        base64data=base64.b64encode(output.getvalue()).decode("ascii"),
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{filename.stem}-replaced.docx",
    )
