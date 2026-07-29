"""python-docx implementation of permanent text data replacement."""

from __future__ import annotations

import base64
import re
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument

from src.core.engines.docx_redaction import _paragraphs
from src.core.models.document import Document
from src.core.models.replacement import ReplacementTarget
from src.core.services.replacement_text import replacement_text


def _replace(paragraph: object, pattern: re.Pattern[str], target: ReplacementTarget) -> None:
    text, count = pattern.subn(lambda match: replacement_text(match.group(), target), paragraph.text)
    if count:
        # Rebuilding the paragraph removes the original runs from the DOCX package.
        paragraph.clear()
        paragraph.add_run(text)


def replace_docx_document(document: Document, targets: list[ReplacementTarget], *, data: bytes | None = None) -> Document:
    """Replace exact text in body, table, header, and footer paragraphs."""
    if data is None and document.base64data is None:
        raise ValueError("DOCX engine requires in-memory document base64data")
    source = DocxDocument(BytesIO(data if data is not None else document.decoded_bytes()))
    paragraphs = _paragraphs(source)
    for target in targets:
        if target.pages is not None:
            raise ValueError("DOCX data replacement does not support page-restricted targets")
        flags = re.IGNORECASE if target.ignore_case else 0
        pattern = re.compile("|".join(re.escape(value) for value in sorted(target.values, key=len, reverse=True)), flags)
        for paragraph in paragraphs:
            _replace(paragraph, pattern, target)
    output = BytesIO()
    source.save(output)
    filename = Path(document.filename or "document.docx")
    return Document(
        base64data=base64.b64encode(output.getvalue()).decode("ascii"),
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{filename.stem}-replaced.docx",
    )
