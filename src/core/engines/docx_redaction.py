"""python-docx implementation of DOCX text redaction."""

from __future__ import annotations

import base64
import re
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from docx.document import Document as DocxFile
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.core.models.document import Document
from src.core.models.redaction import (
    BoundingBoxTarget,
    PageTarget,
    RedactionOptions,
    RedactionTarget,
    RegexTarget,
    TextTarget,
)


def _paragraphs_in_table(table: Table) -> list[Paragraph]:
    paragraphs = []
    for row in table.rows:
        for cell in row.cells:
            paragraphs.extend(cell.paragraphs)
            for nested_table in cell.tables:
                paragraphs.extend(_paragraphs_in_table(nested_table))
    return paragraphs


def _paragraphs(document: DocxFile) -> list[Paragraph]:
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        paragraphs.extend(_paragraphs_in_table(table))
    for section in document.sections:
        for part in (section.header, section.footer):
            paragraphs.extend(part.paragraphs)
            for table in part.tables:
                paragraphs.extend(_paragraphs_in_table(table))
    return paragraphs


def _replacement(text: str, options: RedactionOptions) -> str:
    return "[REDACT]" if options.redaction_type == "mask" else "*" * len(text)


def _replace(paragraph: Paragraph, pattern: re.Pattern[str], options: RedactionOptions, limit: int = 0) -> int:
    text, replacements = pattern.subn(lambda match: _replacement(match.group(), options), paragraph.text, count=limit)
    if replacements:
        paragraph.clear()
        paragraph.add_run(text)
    return replacements


def _redact_text(paragraphs: list[Paragraph], target: TextTarget, options: RedactionOptions) -> None:
    """Compile once and visit each paragraph once for a text target."""
    flags = re.IGNORECASE if target.ignore_case else 0
    pattern = re.compile("|".join(re.escape(value) for value in sorted(target.values, key=len, reverse=True)), flags)
    for paragraph in paragraphs:
        _replace(paragraph, pattern, options)


def _redact_regex(paragraphs: list[Paragraph], target: RegexTarget, options: RedactionOptions) -> None:
    flags = re.IGNORECASE if target.ignore_case else 0
    if not target.allow_unicode:
        flags |= re.ASCII
    patterns = [re.compile(value, flags) for value in target.patterns]
    for pattern in patterns:
        for paragraph in paragraphs:
            replacements = _replace(paragraph, pattern, options, limit=1 if target.only_first_match else 0)
            if replacements and target.only_first_match:
                return


def redact_docx_document(
    document: Document, targets: list[RedactionTarget], options: RedactionOptions, *, data: bytes | None = None
) -> Document:
    """Apply text and regular-expression redactions to a DOCX package."""
    if data is None and document.base64data is None:
        raise ValueError("DOCX engine requires in-memory document base64data")
    source = DocxDocument(BytesIO(data if data is not None else document.decoded_bytes()))
    paragraphs = _paragraphs(source)
    for target in targets:
        if isinstance(target, TextTarget):
            if target.pages is not None:
                raise ValueError("DOCX redaction does not support page-restricted targets")
            _redact_text(paragraphs, target, options)
        elif isinstance(target, RegexTarget):
            if target.pages is not None:
                raise ValueError("DOCX redaction does not support page-restricted targets")
            _redact_regex(paragraphs, target, options)
        elif isinstance(target, (BoundingBoxTarget, PageTarget)):
            raise ValueError("DOCX redaction supports only text and regex targets")
    output = BytesIO()
    source.save(output)
    filename = Path(document.filename or "document.docx")
    return Document(
        base64data=base64.b64encode(output.getvalue()).decode("ascii"),
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{filename.stem}-redacted.docx",
    )
